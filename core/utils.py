## this script contains functions used in main.py

# import packages
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from reportlab.pdfbase.pdfmetrics import stringWidth
from docx import Document
from docx.shared import Pt, Inches
import os
import pdfplumber

# load study_material
def load_study_material(path):
    '''Loads study material from a .docx or .pdf file and returns the text content as a string.'''
    ext = os.path.splitext(path)[1].lower()

    if ext == ".docx":
        doc = Document(path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text

    elif ext == ".pdf":
        text = ""
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()

    else:
        raise ValueError(f"Unsupported file type: {ext}. Only .docx and .pdf are supported.")

def load_txt(filename):
    ##stuk aangepast moet nog reviewen of het daadwerkelijk goed werkt
    '''Loads text from a .txt file and returns the content as a string.'''
    filepath = get_filepath(filename)
    
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()
    
    ## oude def load_txt(filename):
    ##'''Loads text from a .txt file and returns the content as a string.'''
    ##with open(filename, "r", encoding="utf-8") as f:
    ##    return f.read()

# function to save output to pdf
def wrap_text(text, font_name, font_size, max_width):
    '''Wraps text to fit within a specified width. Used to create a nice pdf output.'''
    words = text.split()
    lines = []
    current_line = ""
    for word in words:
        test_line = (current_line + " " + word).strip()
        if stringWidth(test_line, font_name, font_size) <= max_width:
            current_line = test_line
        else:
            lines.append(current_line)
            current_line = word
    if current_line:
        lines.append(current_line)
    return lines

def save_results_to_pdf(filename, results_dict):

    ##stuk aangepast moet nog reviewen door anderen
    filepath = get_filepath(filename)
    
    '''Saves the results in a PDF file with proper formatting.'''
    c = canvas.Canvas(filepath, pagesize=A4)
    width, height = A4

    margin = 2 * cm
    usable_width = width - 2 * margin
    start_y = height - margin
    line_height = 14
    paragraph_spacing = 20  # more space between paragraphs
    font_name = "Helvetica"
    font_size = 10

    y = start_y
    c.setFont(font_name, font_size)

    for header, content in results_dict.items():
        # Add section header
        c.setFont("Helvetica-Bold", 14)
        c.drawString(margin, y, header)
        y -= line_height * 1.5

        c.setFont(font_name, font_size)

        for line in content.splitlines():
            if y < margin:
                c.showPage()
                c.setFont(font_name, font_size)
                y = start_y

            if line.strip() == "":
                y -= paragraph_spacing  # paragraph break
                continue

            wrapped_lines = wrap_text(line, font_name, font_size, usable_width)
            for wline in wrapped_lines:
                if y < margin:
                    c.showPage()
                    c.setFont(font_name, font_size)
                    y = start_y
                c.drawString(margin, y, wline)
                y -= line_height

        y -= paragraph_spacing  # space between sections

    c.save()

def save_results_to_docx(filename, results_dict):
    
    ##stuk aangepast moet nog reviewen of daadwerkelijk goed werkt
    filepath = get_filepath(filename)
    
    '''Saves the results in a DOCX file with proper formatting.'''
    doc = Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for header, content in results_dict.items():
        # Add section header
        header_paragraph = doc.add_paragraph()
        run = header_paragraph.add_run(header)
        run.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph("")  # add spacing after header

        # Add content lines
        for line in content.splitlines():
            if line.strip() == "":
                doc.add_paragraph("")  # blank line for paragraph spacing
                continue
            p = doc.add_paragraph(line)
            p.style.font.name = "Arial"
            p.style.font.size = Pt(10)

        # Add extra space between sections
        doc.add_paragraph("")

    doc.save(filepath)


## haalt de directory van de gebruiker op en geeft deze door
def get_filepath(filename):
    import os
    base_directory = os.path.dirname(os.path.abspath(__file__))
    project_directory = os.path.dirname(base_directory)
    
    filepath = os.path.join(project_directory, filename)
    return filepath
