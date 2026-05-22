from docx import Document
from fpdf import FPDF
import io


def generate_docx(text):
    doc = Document()
    doc.add_heading("Gegenereerd Examen", 0)
    doc.add_paragraph(text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    for line in text.split("\n"):
        pdf.cell(200, 10, txt=line.encode("latin-1", "replace").decode("latin-1"), ln=1)

    pdf_bytes = pdf.output(dest="S").encode("latin-1")
    buffer = io.BytesIO(pdf_bytes)
    buffer.seek(0)
    return buffer